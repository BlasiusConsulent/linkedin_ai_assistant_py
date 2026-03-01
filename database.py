"""
Local Data Storage using Plain Text Files (No Pickle)
Also generates a Word document with all posts
"""

import json
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, Any, List, Optional
import os

# Importa python-docx per generare il documento Word
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  Attenzione: python-docx non installato. Installa con: pip install python-docx")

DATA_DIR = Path("data")
DATA_DIR.mkdir(exist_ok=True)

def load_json_txt(filename: str) -> Dict[Any, Any]:
    """Load dictionary from plain text JSON file"""
    filepath = DATA_DIR / filename
    if filepath.exists():
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if content:
                    return json.loads(content)
                else:
                    return {}
        except Exception as e:
            print(f"Error loading {filename}: {e}")
    return {}

def save_json_txt(filename: str, data: Dict[Any, Any]):
    """Save dictionary to plain text JSON file"""
    filepath = DATA_DIR / filename
    try:
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Error saving {filename}: {e}")

def load_list_txt(filename: str) -> List[Dict[str, Any]]:
    """Load list from plain text JSON file"""
    filepath = DATA_DIR / filename
    if filepath.exists():
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if content:
                    loaded = json.loads(content)
                    return loaded if isinstance(loaded, list) else []
                else:
                    return []
        except Exception as e:
            print(f"Error loading {filename}: {e}")
    return []

def save_list_txt(filename: str, data: List[Dict[str, Any]]):
    """Save list to plain text JSON file"""
    filepath = DATA_DIR / filename
    try:
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Error saving {filename}: {e}")

# Funzioni per dati specifici
def get_news() -> Dict[str, Dict[str, Any]]:
    return load_json_txt("news.txt")

def save_news(news_dict: Dict[str, Dict[str, Any]]):
    save_json_txt("news.txt", news_dict)

def get_posts() -> Dict[str, Dict[str, Any]]:
    return load_json_txt("posts.txt")

def save_posts(posts_dict: Dict[str, Dict[str, Any]]):
    save_json_txt("posts.txt", posts_dict)
    # Aggiorna automaticamente il documento Word
    if DOCX_AVAILABLE:
        update_word_document()

def get_alerts() -> List[Dict[str, Any]]:
    return load_list_txt("alerts.txt")

def save_alerts(alerts_list: List[Dict[str, Any]]):
    save_list_txt("alerts.txt", alerts_list)

def get_analytics() -> Dict[str, Any]:
    return load_json_txt("analytics.txt")

def save_analytics(analytics_dict: Dict[str, Any]):
    save_json_txt("analytics.txt", analytics_dict)

# Funzioni helper
def add_news_item(item: Dict[str, Any]):
    news = get_news()
    news[item['id']] = item
    save_news(news)

def add_post(content: str, news_id: Optional[str] = None):
    posts = get_posts()
    post_id = f"post_{int(datetime.now().timestamp())}"
    posts[post_id] = {
        'content': content,
        'news_id': news_id,
        'created_at': datetime.now().isoformat(),
        'published': False
    }
    save_posts(posts)

def add_alert(alert_type: str, title: str, message: str):
    alerts = get_alerts()
    alerts.append({
        'type': alert_type,
        'title': title,
        'message': message,
        'created_at': datetime.now().isoformat(),
        'read': False
    })
    save_alerts(alerts)

def get_unread_alerts() -> List[Dict[str, Any]]:
    return [a for a in get_alerts() if not a.get('read', False)]

def mark_alert_read(alert_index: int):
    alerts = get_alerts()
    if 0 <= alert_index < len(alerts):
        alerts[alert_index]['read'] = True
        save_alerts(alerts)

def update_word_document():
    """Aggiorna automaticamente il documento Word con tutti i post"""
    if not DOCX_AVAILABLE:
        return
    
    try:
        # Crea o carica il documento
        doc_path = DATA_DIR / "linkedin_posts.docx"
        if doc_path.exists():
            doc = Document(doc_path)
        else:
            doc = Document()
            # Aggiungi titolo
            title = doc.add_heading('LinkedIn AI Assistant - Posts Generati', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # Ottieni tutti i post
        posts = get_posts()
        
        # Ordina i post per data (dal più recente al meno recente)
        sorted_posts = sorted(posts.values(), key=lambda x: x['created_at'], reverse=True)
        
        # Aggiungi i nuovi post al documento
        added_count = 0
        for post in sorted_posts:
            # Controlla se il post è già nel documento (controllo basilare)
            post_exists = False
            for paragraph in doc.paragraphs:
                if post['content'][:100] in paragraph.text:  # Controllo sui primi 100 caratteri
                    post_exists = True
                    break
            
            if not post_exists:
                # Aggiungi intestazione del post
                doc.add_heading(f'Post del {post["created_at"]}', level=1)
                
                # Aggiungi il contenuto del post
                doc.add_paragraph(post['content'])
                
                # Aggiungi separatore
                doc.add_paragraph('-' * 50)
                
                added_count += 1
        
        # Salva il documento
        doc.save(doc_path)
        
        if added_count > 0:
            print(f"📝 Aggiornato documento Word con {added_count} nuovi post: {doc_path}")
    
    except Exception as e:
        print(f"⚠️  Errore aggiornamento documento Word: {e}")

def install_docx():
    """Installa python-docx se non presente"""
    if not DOCX_AVAILABLE:
        print("📦 Installazione python-docx per generazione documenti Word...")
        try:
            import subprocess
            subprocess.check_call(['pip', 'install', 'python-docx'])
            print("✅ python-docx installato correttamente!")
            print("   Riavvia l'applicazione per abilitare la generazione Word.")
        except Exception as e:
            print(f"❌ Impossibile installare python-docx: {e}")
            print("   Esegui manualmente: pip install python-docx")
